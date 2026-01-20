from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Optional
import io
from logger import Logger


class Document:
    """Document representation"""

    def __init__(self, text: str, source: str, metadata: Optional[dict] = None):
        self.text = text
        self.source = source
        self.metadata = metadata or {}


class DocumentLoader(ABC):
    """Abstract document loader interface"""

    @abstractmethod
    def load(self, filepath: Path) -> List[Document]:
        """Load documents from file"""
        pass

    @abstractmethod
    def supports(self, filepath: Path) -> bool:
        """Check if loader supports this file type"""
        pass


class TxtLoader(DocumentLoader):
    """Text file loader"""

    def __init__(self):
        self.logger = Logger().get_logger("TxtLoader")

    def supports(self, filepath: Path) -> bool:
        return filepath.suffix.lower() == ".txt"

    def load(self, filepath: Path) -> List[Document]:
        """Load text file"""
        self.logger.info(f"Loading TXT file: {filepath.name}")

        try:
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(filepath, "r", encoding="latin-1") as f:
                content = f.read()

        if not content.strip():
            self.logger.warning(f"File {filepath.name} is empty")
            return []

        # Split by double newlines for document separation
        doc_texts = [text.strip() for text in content.split("\n\n") if text.strip()]

        documents = [
            Document(
                text=text,
                source=str(filepath),
                metadata={"filename": filepath.name, "format": "txt"},
            )
            for text in doc_texts
        ]

        self.logger.info(f"Loaded {len(documents)} documents from {filepath.name}")
        return documents


class PdfLoader(DocumentLoader):
    """PDF file loader"""

    def __init__(self):
        self.logger = Logger().get_logger("PdfLoader")
        self._check_dependencies()

    def _check_dependencies(self):
        """Check if PDF dependencies are available"""
        try:
            import pypdf

            self.pypdf = pypdf
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF support. Install with: pip install pypdf"
            )

    def supports(self, filepath: Path) -> bool:
        return filepath.suffix.lower() == ".pdf"

    def load(self, filepath: Path) -> List[Document]:
        """Load PDF file"""
        self.logger.info(f"Loading PDF file: {filepath.name}")

        try:
            with open(filepath, "rb") as f:
                pdf_reader = self.pypdf.PdfReader(f)

                if len(pdf_reader.pages) == 0:
                    self.logger.warning(f"PDF {filepath.name} has no pages")
                    return []

                documents = []

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()

                        if text and text.strip():
                            # Split page into paragraphs
                            paragraphs = [
                                p.strip() for p in text.split("\n\n") if p.strip()
                            ]

                            for para_idx, para in enumerate(paragraphs):
                                if len(para) >= 50:  # Minimum paragraph length
                                    documents.append(
                                        Document(
                                            text=para,
                                            source=str(filepath),
                                            metadata={
                                                "filename": filepath.name,
                                                "format": "pdf",
                                                "page": page_num + 1,
                                                "paragraph": para_idx + 1,
                                            },
                                        )
                                    )
                    except Exception as e:
                        self.logger.warning(
                            f"Error extracting page {page_num + 1}: {e}"
                        )
                        continue

                self.logger.info(
                    f"Loaded {len(documents)} documents from {filepath.name}"
                )
                return documents

        except Exception as e:
            self.logger.error(f"Failed to load PDF {filepath.name}: {e}")
            raise RuntimeError(f"PDF loading failed: {e}")


class DocxLoader(DocumentLoader):
    """DOCX file loader"""

    def __init__(self):
        self.logger = Logger().get_logger("DocxLoader")
        self._check_dependencies()

    def _check_dependencies(self):
        """Check if DOCX dependencies are available"""
        try:
            import docx

            self.docx = docx
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX support. Install with: pip install python-docx"
            )

    def supports(self, filepath: Path) -> bool:
        return filepath.suffix.lower() in [".docx", ".doc"]

    def load(self, filepath: Path) -> List[Document]:
        """Load DOCX file"""
        self.logger.info(f"Loading DOCX file: {filepath.name}")

        try:
            doc = self.docx.Document(filepath)

            if not doc.paragraphs:
                self.logger.warning(f"DOCX {filepath.name} has no paragraphs")
                return []

            documents = []

            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()

                if text and len(text) >= 50:  # Minimum paragraph length
                    documents.append(
                        Document(
                            text=text,
                            source=str(filepath),
                            metadata={
                                "filename": filepath.name,
                                "format": "docx",
                                "paragraph": para_idx + 1,
                            },
                        )
                    )

            self.logger.info(f"Loaded {len(documents)} documents from {filepath.name}")
            return documents

        except Exception as e:
            self.logger.error(f"Failed to load DOCX {filepath.name}: {e}")
            raise RuntimeError(f"DOCX loading failed: {e}")


class DocumentLoaderFactory:
    """Factory for creating appropriate document loaders"""

    def __init__(self):
        self.logger = Logger().get_logger("LoaderFactory")
        self.loaders = [TxtLoader(), PdfLoader(), DocxLoader()]

    def get_loader(self, filepath: Path) -> Optional[DocumentLoader]:
        """Get appropriate loader for file"""
        for loader in self.loaders:
            if loader.supports(filepath):
                return loader

        self.logger.warning(f"No loader found for {filepath.suffix}")
        return None

    def load_file(self, filepath: Path) -> List[Document]:
        """Load file using appropriate loader"""
        loader = self.get_loader(filepath)

        if loader is None:
            raise ValueError(
                f"Unsupported file format: {filepath.suffix}\n"
                f"Supported formats: .txt, .pdf, .docx"
            )

        return loader.load(filepath)
